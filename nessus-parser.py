import xml.etree.ElementTree as ET
import sys
import re
import argparse
import docx
from docx.shared import Cm
from docx.enum.style import WD_STYLE

def parse_installed_software(software_text):
    software_list = []
    for line in software_text.split('\n'):
        match = re.search(r'cpe:/.*->\s+(.*)', line)
        if match:
            software = match.group(1)
            software_list.append(software)
    return software_list


def process_microsoft_patches(vulnerabilities, finding_name, description, plugin_output, host):
    if finding_name not in vulnerabilities:
        vulnerabilities[finding_name] = {
            'description': description,
            'plugin_output': plugin_output,
            'targets': []
        }
    if host not in vulnerabilities[finding_name]['targets']:
        vulnerabilities[finding_name]['targets'].append(host)


def process_third_party_vulnerabilities(vulnerabilities, finding_name, description, plugin_output, host):
    if finding_name not in vulnerabilities:
        vulnerabilities[finding_name] = {
            'description': description,
            'plugin_output': plugin_output,
            'targets': []
        }
    if host not in vulnerabilities[finding_name]['targets']:
        vulnerabilities[finding_name]['targets'].append(host)


def parse_nessus_file(file_name, microsoft_patches, third_party):
    tree = ET.parse(file_name)
    root = tree.getroot()

    installed_software = {}
    vulnerabilities = {}
    ms_bulletins_and_kbs = set()

    for host in root.findall(".//ReportHost"):
        target = host.get("name")
        for item in host.findall('.//ReportItem[@pluginID="45590"]'):
            software_text = item.find('plugin_output').text.strip()
            installed_software[target] = parse_installed_software(software_text)

        if microsoft_patches:
            for item in host.findall('.//ReportItem[@pluginID="38153"]'):
                plugin_output = item.find('plugin_output').text.strip()
                for line in plugin_output.split('\n'):
                    match = re.search(r'- (MS\d+-\d+|KB\d+)', line)
                    if match:
                        ms_bulletins_and_kbs.add(match.group(1))
            for finding in root.findall('.//ReportItem'):
                finding_name = finding.find('plugin_name').text
                if finding_name.startswith("Security Update for"):
                    ms_bulletins_and_kbs.add(finding_name)

    for finding in root.findall('.//ReportItem'):
        plugin_id = finding.get("pluginID")
        severity = int(finding.get("severity"))
        if severity >= 1:
            finding_name = finding.find('plugin_name').text
            description = finding.find('description').text.strip()
            plugin_output_element = finding.find('plugin_output')
            plugin_output = plugin_output_element.text.strip() if plugin_output_element is not None else 'N/A'
            target = finding.get('port')

            for host, software_list in installed_software.items():
                combined_software_and_patches = software_list + list(ms_bulletins_and_kbs)
                for item in combined_software_and_patches:
                    if item in finding_name:
                        matched_bulletin_or_kb = any(bulletin_or_kb in finding_name for bulletin_or_kb in ms_bulletins_and_kbs)
                        
                        if microsoft_patches and matched_bulletin_or_kb:
                            process_microsoft_patches(vulnerabilities, finding_name, description, plugin_output, host)
                        elif third_party and not matched_bulletin_or_kb and "Microsoft" not in item:
                            process_third_party_vulnerabilities(vulnerabilities, finding_name, description, plugin_output, host)

    return vulnerabilities
    
def create_styles(doc):
    styles = doc.styles
    
    # Define finding_name style
    finding_name_style = styles.add_style('finding_name', docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
    finding_name_style.font.name = 'Open Sans'
    finding_name_style.font.bold = True
    finding_name_style.font.size = docx.shared.Pt(11)
    finding_name_style.font.color.rgb = docx.shared.RGBColor(74, 193, 224)  # Set font color to #4AC1E0
    finding_name_style.paragraph_format.space_after = docx.shared.Pt(6)
    finding_name_style.paragraph_format.line_spacing_rule = docx.enum.text.WD_LINE_SPACING.AT_LEAST
    finding_name_style.paragraph_format.line_spacing = docx.shared.Pt(14)
    finding_name_style.paragraph_format.keep_together = True
    finding_name_style.paragraph_format.keep_with_next = True

    # Define Details style
    details_style = styles.add_style('details_style', docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
    details_style.font.name = 'Open Sans'
    details_style.font.bold = True
    details_style.font.size = docx.shared.Pt(9.5)
    details_style.paragraph_format.space_after = docx.shared.Pt(6)
    
    # Define Affected Targets style
    affected_targets_style = styles.add_style('affected_targets_style', docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
    affected_targets_style.font.name = 'Open Sans'
    affected_targets_style.font.bold = True
    affected_targets_style.font.size = docx.shared.Pt(9.5)
    affected_targets_style.paragraph_format.space_after = docx.shared.Pt(6)
    
    # Define description style
    description_style = styles.add_style('description', docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
    description_style.font.name = 'Open Sans Light'
    description_style.font.size = docx.shared.Pt(9.5)
    description_style.paragraph_format.space_after = docx.shared.Pt(6)

    # Define plugin_output style
    plugin_output_style = styles.add_style('plugin_output', docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
    plugin_output_style.font.name = 'Courier New'
    plugin_output_style.font.size = docx.shared.Pt(9.5)
    plugin_output_style.paragraph_format.space_after = docx.shared.Pt(6)
    plugin_output_style.paragraph_format.left_indent = docx.shared.Cm(0.2)
    plugin_output_style.paragraph_format.right_indent = docx.shared.Cm(0.2)

    # Define targets style
    targets_style = styles.add_style('targets', docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
    targets_style.font.name = 'Open Sans Light'
    targets_style.font.size = docx.shared.Pt(9.5)
    targets_style.paragraph_format.space_after = docx.shared.Pt(6)

    # Define bullet style
    bullet_style = styles.add_style('bullet_style', docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
    bullet_style.font.name = 'Open Sans Light'
    bullet_style.font.size = docx.shared.Pt(9.5)
    bullet_style.paragraph_format.left_indent = docx.shared.Cm(0)
    bullet_style.paragraph_format.first_line_indent = docx.shared.Cm(-0.6)
    
    
def print_output(vulnerabilities, output_format="docx", output_file="output.docx"):
    if output_format == "docx":
        doc = docx.Document()

        # Set the page layout to landscape
        section = doc.sections[0]
        section.page_height, section.page_width = section.page_width, section.page_height

        # Set the margins
        section.left_margin = Cm(1.5)
        section.top_margin = Cm(2)
        section.right_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.gutter = Cm(0)

        create_styles(doc)

        for finding_name, info in vulnerabilities.items():
            p_finding_name = doc.add_paragraph(style='finding_name')
            p_finding_name.add_run(finding_name)

            p_description = doc.add_paragraph(style='description')
            description_text = info['description'].split('\n\n', 1)[0]
            
            #Legacy Code - Keeping it in here for reference
            #description_text = info['description'].replace("  - ", "\n• ")
            
            description_lines = description_text.split('\n')
            
            #Below is the bullet point magic
            for line in description_lines:
                if line.startswith("• "):
                    bullet_paragraph = doc.add_paragraph(style='bullet_style')
                    bullet_paragraph.add_run(line[2:])
                    bullet_paragraph.paragraph_format.first_line_indent = docx.shared.Cm(-0.6)
                    bullet_paragraph.paragraph_format.left_indent = docx.shared.Cm(0.6)
                    bullet_paragraph.paragraph_format.space_after = docx.shared.Pt(0)
                    bullet_paragraph.paragraph_format.tab_stops.add_tab_stop(Cm(0.6))

                    num_pr = bullet_paragraph._element.get_or_add_pPr().get_or_add_numPr()
                    num_pr.get_or_add_ilvl().set(docx.oxml.ns.qn('w:val'), "0")
                    num_pr.get_or_add_numId().set(docx.oxml.ns.qn('w:val'), "1")
                else:
                    p_description.add_run(line)
                    p_description.add_run(' ')
                    
            p_details_style_output = doc.add_paragraph(style='details_style')
            p_details_style_output.add_run("Details:")
            
            p_plugin_output = doc.add_paragraph(style='plugin_output')
            p_plugin_output.add_run(info['plugin_output'])

            p_targets = doc.add_paragraph(style='targets')
            p_affected_targets_style_output = doc.add_paragraph(style='affected_targets_style')
            p_affected_targets_style_output.add_run("Affected Targets:")

            for target in info['targets']:
                p_bullet = doc.add_paragraph(style='bullet_style')
                p_bullet.add_run(target)
                p_bullet.paragraph_format.first_line_indent = docx.shared.Cm(-0.6)
                p_bullet.paragraph_format.left_indent = docx.shared.Cm(0.6)

        doc.save("output.docx")
    elif output_format == "cli":
        for finding_name, info in vulnerabilities.items():
            print(f"\n{finding_name}\n")
            print("Description:")
            print(info['description'])
            print("\nDetails:")
            print(info['plugin_output'])
            print("\nAffected Targets:")
            for target in info['targets']:
                print(f"  - {target}")

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Nessus parser for extracting outdated software and associated targets")
    parser.add_argument("-f", "--file", dest="file_name", required=True, help="Path to the Nessus (.nessus) file")
    parser.add_argument("-mp", "--microsoft-patches", action="store_true", help="Only include findings related to Microsoft missing patches")
    parser.add_argument("-tp", "--third-party", action="store_true", help="Only include findings related to third-party outdated software")
    parser.add_argument("-o", "--output", dest="output_file", default="output.docx", help="Output file name (default: output.docx)")
    parser.add_argument("-fmt", "--format", dest="output_format", choices=["docx", "cli"], default="docx", help="Output format: docx (Word document) or cli (command-line) (default: docx)")


    args = parser.parse_args()

    vulnerabilities = parse_nessus_file(args.file_name, args.microsoft_patches, args.third_party)
    print_output(vulnerabilities, args.output_format, args.output_file)