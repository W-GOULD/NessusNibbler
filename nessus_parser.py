import xml.etree.ElementTree as ET
import sys
import os
import re
import argparse
import docx
from docx.shared import Cm
from docx.enum.style import WD_STYLE
from styles import create_styles


def parse_installed_software(software_text):
    software_list = []
    for line in software_text.split('\n'):
        match = re.search(r'cpe:/.*->\s+(.*)', line)
        if match:
            software = match.group(1)
            software_list.append(software)
    return software_list


def process_microsoft_patches(vulnerabilities, finding_name, description, plugin_output, host):
    if finding_name not in vulnerabilities:
        vulnerabilities[finding_name] = {
            'description': description,
            'plugin_output': plugin_output,
            'targets': []
        }
    if host not in vulnerabilities[finding_name]['targets']:
        vulnerabilities[finding_name]['targets'].append(host)

def process_third_party_vulnerabilities(vulnerabilities, finding_name, description, plugin_output, host):
    if finding_name not in vulnerabilities:
        vulnerabilities[finding_name] = {
            'description': description,
            'plugin_output': plugin_output,
            'targets': []
        }
    if host not in vulnerabilities[finding_name]['targets']:
        vulnerabilities[finding_name]['targets'].append(host)

def process_unquoted_service_path_vulnerabilities(vulnerabilities, finding_name, description, plugin_output, host):
    if finding_name not in vulnerabilities:
        vulnerabilities[finding_name] = {
            'description': description,
            'plugin_output': {},
            'targets': []
        }
    if host not in vulnerabilities[finding_name]['targets']:
        vulnerabilities[finding_name]['targets'].append(host)
    vulnerabilities[finding_name]['plugin_output'][host] = plugin_output

def clean_description_text(description_text):
    cleaned_text = re.sub(r'<code>|</code>', '', description_text)
    cleaned_text = re.sub(r' {2,}', ' ', cleaned_text)
    return cleaned_text

def parse_nessus_file(file_name, microsoft_patches, third_party, unquoted_service_path):
    tree = ET.parse(file_name)
    root = tree.getroot()

    installed_software = {}
    vulnerabilities = {}
    ms_bulletins_and_kbs = set()

    for host in root.findall(".//ReportHost"):
        target = host.get("name")
        for item in host.findall('.//ReportItem[@pluginID="45590"]'):
            software_text = item.find('plugin_output').text.strip()
            installed_software[target] = parse_installed_software(software_text)

        if microsoft_patches:
            for item in host.findall('.//ReportItem[@pluginID="38153"]'):
                plugin_output = item.find('plugin_output').text.strip()
                for line in plugin_output.split('\n'):
                    match = re.search(r'- (MS\d+-\d+|KB\d+)', line)
                    if match:
                        ms_bulletins_and_kbs.add(match.group(1))
            for finding in root.findall('.//ReportItem'):
                finding_name = finding.find('plugin_name').text
                if finding_name.startswith("Security Update for"):
                    ms_bulletins_and_kbs.add(finding_name)

        for finding in host.findall('.//ReportItem'):
            plugin_id = finding.get("pluginID")
            severity = int(finding.get("severity"))
            if severity >= 1:
                finding_name = finding.find('plugin_name').text
                description = clean_description_text(finding.find('description').text.strip())
                plugin_output_element = finding.find('plugin_output')
                plugin_output = plugin_output_element.text.strip() if plugin_output_element is not None else 'N/A'

                if plugin_id == "63155" and unquoted_service_path:  # Microsoft Windows Unquoted Service Path Enumeration
                    process_unquoted_service_path_vulnerabilities(vulnerabilities, finding_name, description, plugin_output, target)
                else:
                    for host, software_list in installed_software.items():
                        combined_software_and_patches = software_list + list(ms_bulletins_and_kbs)
                        for item in combined_software_and_patches:
                            if item in finding_name:
                                matched_bulletin_or_kb = any(bulletin_or_kb in finding_name for bulletin_or_kb in ms_bulletins_and_kbs)
                                if microsoft_patches and matched_bulletin_or_kb:
                                    process_microsoft_patches(vulnerabilities, finding_name, description, plugin_output, host)
                                elif third_party and not matched_bulletin_or_kb and "Microsoft" not in item:
                                    process_third_party_vulnerabilities(vulnerabilities, finding_name, description, plugin_output, host)
    return vulnerabilities

def explore_nessus_file(file_name):
    if file_name is None:
        return [], {}

    tree = ET.parse(file_name)
    root = tree.getroot()

    findings = []

    for host in root.findall(".//ReportHost"):
        target = host.get("name")
        os_tag = host.find('./HostProperties/tag[@name="os"]')
        os = os_tag.text if os_tag is not None else "Unknown"
        hostname = host.find('./HostProperties/tag[@name="host-fqdn"]').text if host.find('./HostProperties/tag[@name="host-fqdn"]') is not None else 'N/A'

        for item in host.findall('.//ReportItem'):
            plugin_id = item.get("pluginID")
            plugin_name = item.find('plugin_name').text
            severity = int(item.get("severity"))
            port = item.get("port")
            service = item.get("svc_name")

            if severity >= 0:
                description = clean_description_text(item.find('description').text.strip())
                synopsis = item.find('synopsis').text.strip()
                solution = item.find('solution').text.strip()
                plugin_output_element = item.find('plugin_output')
                plugin_output = plugin_output_element.text.strip() if plugin_output_element is not None and plugin_output_element.text is not None else 'N/A'
                cve = item.find('cve').text.strip() if item.find('cve') is not None else 'N/A'
                cvss3_base_score = item.find('cvss3_base_score').text.strip() if item.find('cvss3_base_score') is not None else 'N/A'
                cvss3_vector = item.find('cvss3_vector').text.strip() if item.find('cvss3_vector') is not None else 'N/A'
                external_reference = item.find('see_also').text.strip() if item.find('see_also') is not None else 'N/A'

                finding = {
                    'id': plugin_id,
                    'plugin_name': plugin_name,
                    'host_ip': target,
                    'port': port,
                    'service': service,
                    'os': os,
                    'hostname': hostname,
                    'risk_rating': severity,
                    'external_reference': external_reference,
                    'cvssv3': cvss3_base_score,
                    'cvssv3_vector': cvss3_vector,
                    'description': description,
                    'synopsis': synopsis,
                    'solution': solution,
                    'plugin_output': plugin_output,
                    'cve': cve
                }

                findings.append(finding)

    return findings

def print_output(vulnerabilities, output_format="docx", output_file="output.docx"):
    if output_format == "docx":
        doc = docx.Document()

        # Set the page layout to landscape
        section = doc.sections[0]
        section.page_height, section.page_width = section.page_width, section.page_height

        # Set the margins
        section.left_margin = Cm(1.5)
        section.top_margin = Cm(2)
        section.right_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.gutter = Cm(0)

        create_styles(doc)
        for finding_name, info in vulnerabilities.items():
            if finding_name == "Microsoft Windows Unquoted Service Path Enumeration":
                p_finding_name = doc.add_paragraph(style='finding_name')
                p_finding_name.add_run(finding_name)

                p_description = doc.add_paragraph(style='description')
                description_text = info['description'].split('\n\n', 1)[0]
                p_description.add_run(description_text)

                for target in info['targets']:
                    p_targets = doc.add_paragraph(style='targets')
                    p_targets.add_run(target)

                    p_plugin_output = doc.add_paragraph(style='plugin_output')
                    p_plugin_output.add_run(info['plugin_output'][target])
            else:
                for finding_name, info in vulnerabilities.items():
                    p_finding_name = doc.add_paragraph(style='finding_name')
                    p_finding_name.add_run(finding_name)

                    p_description = doc.add_paragraph(style='description')
                    description_text = info['description'].split('\n\n', 1)[0]
                    description_lines = description_text.split('\n')
                    
                    #Below is the bullet point magic
                    for line in description_lines:
                        if line.startswith("• "):
                            bullet_paragraph = doc.add_paragraph(style='bullet_style')
                            bullet_paragraph.add_run(line[2:])
                            bullet_paragraph.paragraph_format.first_line_indent = docx.shared.Cm(-0.6)
                            bullet_paragraph.paragraph_format.left_indent = docx.shared.Cm(0.6)
                            bullet_paragraph.paragraph_format.space_after = docx.shared.Pt(0)
                            bullet_paragraph.paragraph_format.tab_stops.add_tab_stop(Cm(0.6))

                            num_pr = bullet_paragraph._element.get_or_add_pPr().get_or_add_numPr()
                            num_pr.get_or_add_ilvl().set(docx.oxml.ns.qn('w:val'), "0")
                            num_pr.get_or_add_numId().set(docx.oxml.ns.qn('w:val'), "1")
                        else:
                            p_description.add_run(line)
                            p_description.add_run(' ')
                            
                    p_details_style_output = doc.add_paragraph(style='details_style')
                    p_details_style_output.add_run("Details:")
                    
                    p_plugin_output = doc.add_paragraph(style='plugin_output')
                    p_plugin_output.add_run(info['plugin_output'])

                    p_targets = doc.add_paragraph(style='targets')
                    p_affected_targets_style_output = doc.add_paragraph(style='affected_targets_style')
                    p_affected_targets_style_output.add_run("Affected Targets:")

                    for target in info['targets']:
                        p_bullet = doc.add_paragraph(style='bullet_style')
                        p_bullet.add_run(target)
                        p_bullet.paragraph_format.first_line_indent = docx.shared.Cm(-0.6)
                        p_bullet.paragraph_format.left_indent = docx.shared.Cm(0.6)

        doc.save(output_file)

    elif output_format == "txt":
        with open(output_file, 'w') as output_txt:
            for finding_name, info in vulnerabilities.items():
                if finding_name == "Microsoft Windows Unquoted Service Path Enumeration":
                    output_txt.write(f"\n{finding_name}\n")
                    description_text = info['description'].split('\n\n', 1)[0]
                    output_txt.write(description_text)
                    output_txt.write("\n\n")

                    for target in info['targets']:
                        output_txt.write(f"{target}\n\n")
                        output_txt.write(info['plugin_output'][target])
                        output_txt.write("\n\n")
                else:
                    output_txt.write(f"\n{finding_name}\n")
                    output_txt.write("Description:\n")
                    description_text = info['description'].split('\n\n', 1)[0]
                    output_txt.write(description_text)
                    output_txt.write("\n\nDetails:\n")
                    output_txt.write(info['plugin_output'])
                    output_txt.write("\n\nAffected Targets:\n")
                    for target in info['targets']:
                        output_txt.write(f"  - {target}\n")

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Nessus parser for extracting outdated software and associated targets")
    parser.add_argument("-f", "--file", dest="file_name", required=True, help="Path to the Nessus (.nessus) file")
    parser.add_argument("-mp", "--microsoft-patches", action="store_true", help="Only include findings related to Microsoft missing patches")
    parser.add_argument("-tp", "--third-party", action="store_true", help="Only include findings related to third-party outdated software")
    parser.add_argument("-o", "--output", dest="output_file", default="output.docx", help="Output file name (default: output.docx)")
    parser.add_argument("-fmt", "--format", dest="output_format", choices=["docx", "txt"], default="docx", help="Output format: docx (Word document) or txt (text file) (default: docx)")
    parser.add_argument("-u", "--unquoted-service-path", action="store_true", help="Only include findings related to unquoted service path vulnerabilities")

    args = parser.parse_args()

    if args.explore:
        hosts, vulnerabilities = explore_nessus_file(args.file_name)
        print_output(vulnerabilities, output_format=args.output_format, output_file=args.output_file)
    else:
        vulnerabilities = parse_nessus_file(args.file_name, args.microsoft_patches, args.third_party)
        print_output(vulnerabilities, output_format=args.output_format, output_file=args.output_file)