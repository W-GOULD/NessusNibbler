from utils.extract import *
import xml.etree.ElementTree as ET
import docx
import csv
from docx.shared import Cm
from docx.enum.style import WD_STYLE
from styles import create_styles
from collections import defaultdict

def write_to_csv(filename, data):
    with open(filename, 'w', newline='') as file:
        writer = csv.writer(file)
        writer.writerow(["Item No", "Benchmark", "Rationale", "Recommendation", "Current Setting", "Targets"])
        writer.writerows(data)


def print_output(vulnerabilities, output_format="docx", output_file="output.docx"):
    if output_format == "docx":
        doc = docx.Document()

        # Set the page layout to landscape
        section = doc.sections[0]
        section.page_height, section.page_width = section.page_width, section.page_height

        # Set the margins
        section.left_margin = Cm(1.5)
        section.top_margin = Cm(2)
        section.right_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.gutter = Cm(0)

        create_styles(doc)
        for finding_name, info in vulnerabilities.items():
            if finding_name == "Microsoft Windows Unquoted Service Path Enumeration":
                p_finding_name = doc.add_paragraph(style='finding_name')
                p_finding_name.add_run(finding_name)

                p_description = doc.add_paragraph(style='description')
                description_text = info['description'].split('\n\n', 1)[0]
                p_description.add_run(description_text)

                for target in info['targets']:
                    p_targets = doc.add_paragraph(style='targets')
                    p_targets.add_run(target)

                    p_plugin_output = doc.add_paragraph(style='plugin_output')
                    p_plugin_output.add_run(info['plugin_output'][target])
            else:
                
                    p_finding_name = doc.add_paragraph(style='finding_name')
                    p_finding_name.add_run(finding_name)

                    p_description = doc.add_paragraph(style='description')
                    description_text = info['description'].split('\n\n', 1)[0]
                    description_lines = description_text.split('\n')
                    
                    #Below is the bullet point magic
                    for line in description_lines:
                        if line.startswith("• "):
                            bullet_paragraph = doc.add_paragraph(style='bullet_style')
                            bullet_paragraph.add_run(line[2:])
                            bullet_paragraph.paragraph_format.first_line_indent = docx.shared.Cm(-0.6)
                            bullet_paragraph.paragraph_format.left_indent = docx.shared.Cm(0.6)
                            bullet_paragraph.paragraph_format.space_after = docx.shared.Pt(0)
                            bullet_paragraph.paragraph_format.tab_stops.add_tab_stop(Cm(0.6))

                            num_pr = bullet_paragraph._element.get_or_add_pPr().get_or_add_numPr()
                            num_pr.get_or_add_ilvl().set(docx.oxml.ns.qn('w:val'), "0")
                            num_pr.get_or_add_numId().set(docx.oxml.ns.qn('w:val'), "1")
                        else:
                            p_description.add_run(line)
                            p_description.add_run(' ')
                            
                    p_details_style_output = doc.add_paragraph(style='details_style')
                    p_details_style_output.add_run("Details:")
                    
                    p_plugin_output = doc.add_paragraph(style='plugin_output')
                    p_plugin_output.add_run(info['plugin_output'])

                    p_targets = doc.add_paragraph(style='targets')
                    p_affected_targets_style_output = doc.add_paragraph(style='affected_targets_style')
                    p_affected_targets_style_output.add_run("Affected Targets:")

                    for target in info['targets']:
                        p_bullet = doc.add_paragraph(style='bullet_style')
                        p_bullet.add_run(target)
                        p_bullet.paragraph_format.first_line_indent = docx.shared.Cm(-0.6)
                        p_bullet.paragraph_format.left_indent = docx.shared.Cm(0.6)

        doc.save(output_file)

    elif output_format == "txt":
        with open(output_file, 'w') as output_txt:
            for finding_name, info in vulnerabilities.items():
                if finding_name == "Microsoft Windows Unquoted Service Path Enumeration":
                    output_txt.write(f"\n{finding_name}\n")
                    description_text = info['description'].split('\n\n', 1)[0]
                    output_txt.write(description_text)
                    output_txt.write("\n\n")

                    for target in info['targets']:
                        output_txt.write(f"{target}\n\n")
                        output_txt.write(info['plugin_output'][target])
                        output_txt.write("\n\n")
                else:
                    output_txt.write(f"\n{finding_name}\n")
                    output_txt.write("Description:\n")
                    description_text = info['description'].split('\n\n', 1)[0]
                    output_txt.write(description_text)
                    output_txt.write("\n\nDetails:\n")
                    output_txt.write(info['plugin_output'])
                    output_txt.write("\n\nAffected Targets:\n")
                    for target in info['targets']:
                        output_txt.write(f"  - {target}\n")


    elif output_format == "csv":
        with open(output_file, 'w', newline='') as file:
            writer = csv.writer(file)
            writer.writerow(["Item No", "Benchmark", "Rationale", "Recommendation", "Current Setting", "Targets"])
            for item, hosts in vulnerabilities.items():
                row = list(item) + ['\n'.join(hosts)]
                writer.writerow(row)
